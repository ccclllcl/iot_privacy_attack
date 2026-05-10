#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""用最终实验结果重写论文第四章，并插入论文专用图。"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
SUMMARY_DIR = ROOT / "outputs" / "summaries" / "final_thesis"
FIG_DIR = ROOT / "outputs" / "figures" / "summaries" / "final_thesis"
DEFAULT_OUTPUT_NAME = "毕业论文_第四章图表与结果最终重写版.docx"


FIGURES = {
    "mock_accuracy": FIG_DIR / "thesis_fig4_01_mock_accuracy.png",
    "mock_distortion": FIG_DIR / "thesis_fig4_02_mock_distortion.png",
    "ldp_scan": FIG_DIR / "thesis_fig4_03_ldp_parameter_scan.png",
    "noise_scan": FIG_DIR / "thesis_fig4_04_noise_parameter_scan.png",
    "adaptive_scan": FIG_DIR / "thesis_fig4_05_adaptive_ldp_parameter_scan.png",
    "confusion_baseline": FIG_DIR / "thesis_fig4_07_confusion_mock_baseline.png",
    "confusion_adaptive": FIG_DIR / "thesis_fig4_08_confusion_mock_adaptive_lstm_fixed.png",
    "real_accuracy": FIG_DIR / "thesis_fig4_10_real_dataset_accuracy.png",
    "cooja_accuracy": FIG_DIR / "thesis_fig4_12_cooja_accuracy.png",
}


def set_run_font(run, size_pt: float = 12.0, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph, text: str, size_pt: float = 12.0) -> None:
    for child in list(paragraph._p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    set_run_font(run, size_pt=size_pt)


def add_before(anchor, text: str = "", style: str | None = None, size_pt: float = 12.0):
    paragraph = anchor.insert_paragraph_before("", style=style)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size_pt=size_pt)
    return paragraph


def add_body(anchor, text: str):
    paragraph = add_before(anchor, text, "Normal", 12.0)
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.space_after = Pt(0)
    return paragraph


def add_heading(anchor, text: str, style: str):
    return add_before(anchor, text, style, 12.0)


def add_caption(anchor, text: str):
    paragraph = add_before(anchor, text, "Normal", 10.5)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_figure(anchor, image_path: Path, caption: str, width_cm: float = 14.2):
    paragraph = add_before(anchor, "", "Normal", 12.0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_caption(anchor, caption)


def find_body_heading(doc: Document, text: str, style_name: str = "Heading 1"):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text and paragraph.style.name == style_name:
            return paragraph
    raise ValueError(f"未找到标题：{text}")


def remove_chapter4(doc: Document):
    start = find_body_heading(doc, "4 实验设计与结果分析", "Heading 1")
    end = find_body_heading(doc, "5 总结与展望", "Heading 1")
    body = doc.element.body
    children = list(body)
    start_idx = children.index(start._p)
    end_idx = children.index(end._p)
    removed = 0
    for element in children[start_idx:end_idx]:
        body.remove(element)
        removed += 1
    return end, removed


def add_update_fields_setting(doc: Document) -> None:
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)
    else:
        existing.set(qn("w:val"), "true")


def replace_outdated_phrases(doc: Document) -> list[dict[str, str]]:
    replacements = {
        "run_compare.py 只对 ldp 和 noise 执行，代表 LSTM + fixed_attacker。": (
            "最终实验结果已经覆盖 adaptive_ldp、ldp、noise 三类方法在 mock 与真实数据上的参数扫描，"
            "并同时包含 LSTM/MLP、fixed_attacker/retrain_attacker 与 3 个 seed。"
        ),
        "真实数据参数扫描主要补齐 UCI HAR": "真实数据参数扫描覆盖 UCI HAR、Kasteren 与 CASAS 三个数据集",
        "后续可以开展 adaptive_ldp 消融实验": "当前结果已包含 adaptive_ldp 的 profile 级消融汇总，后续可继续开展更细粒度真实部署消融",
        "outputs/reports/final_thesis": "outputs/summaries/final_thesis",
        "outputs/reports/full_multiseed_summary.json": "outputs/summaries/final_thesis/final_summary.json",
    }
    changes: list[dict[str, str]] = []
    for paragraph in doc.paragraphs:
        original = paragraph.text
        updated = original
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != original:
            set_paragraph_text(paragraph, updated)
            changes.append({"old": original, "new": updated})
    return changes


def insert_chapter4(anchor) -> None:
    add_heading(anchor, "4 实验设计与结果分析", "Heading 1")

    add_heading(anchor, "4.1 实验环境与结果来源", "Heading 2")
    add_body(
        anchor,
        "本章基于最终统一汇总的实验结果，对 mock 合成智能家居数据、UCI HAR、Kasteren、CASAS 真实公开数据以及 Cooja 节点级日志实验进行分析。"
        "实验不再沿用早期结果未补齐时的单一口径，而是以完整矩阵和参数扫描结果为依据。最终覆盖审计显示，mock 主矩阵完成 36/36，真实数据主矩阵完成 108/108；"
        "mock 参数扫描完成 36/36，真实数据参数扫描完成 108/108；adaptive_ldp 在每个 dataset、seed、model、mode 组合下均包含 6 个 profile；Cooja canonical 结果完成 18/18。",
    )
    add_body(
        anchor,
        "本章重点回答三个问题：第一，设备状态序列和流量统计特征是否足以支持行为推断攻击；第二，noise、ldp、adaptive_ldp 三类防御在 fixed_attacker 与 retrain_attacker 两种威胁模型下如何影响攻击准确率；"
        "第三，隐私抑制与数据失真之间是否存在可解释的权衡关系。由于不同数据集的类别数量、采样方式和任务定义并不一致，真实数据结果只在各自数据集内部比较 baseline 与 defended 变化，不对不同数据集做绝对排名。",
    )
    add_body(
        anchor,
        "攻击模型方面，LSTM 用于学习窗口序列中的时间依赖，MLP 用于学习窗口级统计特征。两者的差异有助于区分“时序结构泄露”和“统计特征泄露”两种风险来源。"
        "防御方法方面，noise 提供加性随机扰动基线，ldp 表示固定隐私预算下的本地差分隐私扰动，adaptive_ldp 则根据窗口波动和流量强度等代理指标动态调整隐私预算。",
    )

    add_heading(anchor, "4.2 攻击基线结果", "Heading 2")
    add_body(
        anchor,
        "mock 主实验首先体现出攻击基线差异。跨 3 个 seed 平均后，LSTM 的 baseline_acc 为 0.6402，MLP 的 baseline_acc 为 0.4735。"
        "这说明行为识别不仅依赖窗口内部的设备触发次数，还明显依赖设备状态变化的先后顺序。LSTM 能够利用“移动触发—设备开启—活动持续”等时间结构，因而在基线阶段表现出更强的行为推断能力。",
    )
    add_body(
        anchor,
        "从隐私风险角度看，基线准确率越高，说明原始数据表示中保留的行为信息越容易被攻击模型学习。MLP 虽然不直接建模时间顺序，但仍能通过统计特征获得高于随机猜测的识别能力，表明窗口级聚合特征本身也可能泄露行为模式。",
    )
    add_figure(
        anchor,
        FIGURES["mock_accuracy"],
        "图4.1 mock 场景下 LSTM/MLP 的 baseline、fixed_attacker 与 retrain_attacker 准确率对比",
    )

    add_heading(anchor, "4.3 多方法防御结果比较", "Heading 2")
    add_heading(anchor, "4.3.1 LSTM 结果分析", "Heading 3")
    add_body(
        anchor,
        "在 LSTM 攻击模型下，三类防御使用相同的基线结果，因此可以直接比较 defended_acc 的变化。默认参数下，ldp 对攻击抑制最强："
        "fixed_attacker 下平均 defended_acc 为 0.3377，retrain_attacker 下为 0.2980。noise 的 fixed_attacker 与 retrain_attacker 平均 defended_acc 分别为 0.3714 和 0.3609；"
        "adaptive_ldp 分别为 0.4542 和 0.4360。由此可见，固定预算 LDP 在默认参数下提供了最强的攻击抑制，但这一结果也伴随更高的数据失真。",
    )
    add_body(
        anchor,
        "retrain_attacker 结果说明，攻击者在防御后数据上重新训练并不能完全恢复原始识别能力。尤其在 LSTM 条件下，扰动破坏了跨时间步的稳定依赖，重训练攻击者仍然难以恢复基线水平。"
        "这表明防御并非只造成模型分布偏移，而是改变了攻击模型可利用的时间结构。",
    )

    add_heading(anchor, "4.3.2 MLP 结果分析", "Heading 3")
    add_body(
        anchor,
        "MLP 的结果与 LSTM 呈现出不同特征。由于 MLP 主要依赖窗口统计量，其基线准确率低于 LSTM，但在 retrain_attacker 下对部分防御存在更明显恢复。"
        "例如 adaptive_ldp 在 MLP fixed_attacker 下平均 defended_acc 为 0.2191，而 retrain_attacker 下回升到 0.3968。"
        "这说明当攻击者只需要重新学习被扰动后的统计边界时，残余统计结构仍可能被利用。",
    )
    add_body(
        anchor,
        "综合 LSTM 与 MLP 结果可见，防御效果与攻击模型能力密切相关。LSTM 更容易暴露时序泄露风险，也更能体现扰动对行为过程结构的破坏；"
        "MLP 则揭示了统计特征在重训练条件下可能保留的剩余风险。因此，评价物联网隐私防御时需要同时考虑固定攻击者和自适应攻击者。",
    )
    add_figure(anchor, FIGURES["mock_distortion"], "图4.2 三种防御方法的 MSE、MAE 与 Pearson_r 对比")

    add_heading(anchor, "4.3.3 参数扫描结果", "Heading 3")
    add_body(
        anchor,
        "最终结果已经补齐 adaptive_ldp、ldp、noise 三类方法在 mock 与真实数据上的参数扫描，覆盖 LSTM/MLP、fixed_attacker/retrain_attacker 与 3 个 seed。"
        "正文选取 mock 场景下 LSTM fixed_attacker 作为代表性曲线，以便清晰展示参数变化方向；完整扫描结果用于支持趋势分析。",
    )
    add_body(
        anchor,
        "LDP 的 epsilon 越小，扰动越强。图4.3显示，当 epsilon 从 0.1 增大到 5.0 时，defended_acc 从 0.2721 上升到 0.5679，MSE 从 99.7281 下降到 0.0421。"
        "这说明更强隐私预算约束能够显著降低攻击准确率，但也会带来非常高的数据失真；反之，epsilon 增大后数据保真度恢复，但剩余识别风险也上升。",
    )
    add_figure(anchor, FIGURES["ldp_scan"], "图4.3 LDP 参数扫描下 defended_acc 与 MSE 的变化趋势")
    add_body(
        anchor,
        "noise 方法呈现相反方向的参数含义。图4.4显示，当 noise_scale 从 0.1 增大到 1.0 时，defended_acc 从 0.6639 下降到 0.3168，MSE 从 0.0200 上升到 1.9962。"
        "因此，噪声强度越高，攻击抑制越明显，但对原始数据结构的破坏也越大。",
    )
    add_figure(anchor, FIGURES["noise_scan"], "图4.4 Noise 参数扫描下 defended_acc 与 MSE 的变化趋势")
    add_body(
        anchor,
        "adaptive_ldp 不是单一连续参数，而是通过 profile 组合改变 epsilon_min、epsilon_max、weight_sensitivity、weight_traffic 和 use_edge_budget_cap。"
        "6 个 profile 分别对应默认配置、强隐私、弱隐私、仅窗口波动、仅流量强度和启用边缘预算裁剪接口。图4.5表明，adaptive_strong_privacy 的 MSE 最高、攻击抑制更强，"
        "adaptive_weak_privacy 的 defended_acc 较高而 MSE 较低，体现出 profile 级隐私—可用性差异。该结果是经验性 profile 扫描观察，不构成形式化理论证明。",
    )
    add_figure(anchor, FIGURES["adaptive_scan"], "图4.5 adaptive_ldp profile 参数扫描下 defended_acc 与 MSE 的变化趋势")

    add_heading(anchor, "4.4 隐私—可用性权衡分析", "Heading 2")
    add_body(
        anchor,
        "隐私保护效果不能只用攻击准确率下降来评价，还需要结合数据失真指标分析。图4.2展示了三类方法默认参数下的 MSE、MAE 与 Pearson_r。"
        "noise 的 MSE 最低且 Pearson_r 最高，说明它在默认强度下保留了更多原始相关结构；但正因如此，剩余识别风险也更高。ldp 的 MSE 和 MAE 更高、Pearson_r 更低，对攻击的抑制更强，但可用性损失更大。",
    )
    add_body(
        anchor,
        "adaptive_ldp 位于两者之间：它通过窗口风险代理动态调整隐私预算，试图避免对全部窗口施加同等强度扰动。默认 profile 下，它在保留部分结构信息的同时降低攻击准确率；"
        "但在不同攻击模型和 retrain_attacker 条件下表现存在差异，说明自适应预算机制仍需要结合具体数据模态和攻击者能力进行参数选择。",
    )
    add_body(
        anchor,
        "因此，本章得到的核心认识是：物联网行为隐私保护是一个多目标权衡问题。攻击准确率下降代表隐私风险降低，MSE、MAE 和 Pearson_r 代表数据结构保留程度。"
        "不同防御方法不是简单优劣关系，而是对应不同应用偏好：强隐私场景更倾向于较强 LDP，保真需求更高的场景可考虑较弱噪声或 adaptive_ldp profile。",
    )

    add_heading(anchor, "4.5 混淆矩阵与错误类型分析", "Heading 2")
    add_body(
        anchor,
        "混淆矩阵可以进一步揭示防御后错误迁移的具体方向。LSTM 基线并不是随机猜测，而是对部分强模式类别具有明显识别能力，同时在语义相近或时序模式相似的类别之间产生误分。"
        "这说明原始序列中包含可被模型捕捉的行为结构。",
    )
    add_figure(anchor, FIGURES["confusion_baseline"], "图4.6 seed 42 下 LSTM 基线混淆矩阵")
    add_body(
        anchor,
        "在 adaptive_ldp + LSTM + fixed_attacker 条件下，预测分布发生明显扩散，多个类别之间的边界被扰动破坏。"
        "这类变化说明防御并不是简单降低一个全局准确率，而是改变了攻击模型对行为类别边界的理解。",
    )
    add_figure(anchor, FIGURES["confusion_adaptive"], "图4.7 seed 42 下 adaptive_ldp 的 LSTM fixed_attacker 混淆矩阵")
    add_body(
        anchor,
        "从错误类型看，防御后的误分具有结构性。相近行为之间的混淆、强特征类别向粗粒度类别坍塌、重训练攻击者对残余结构的利用，"
        "共同说明隐私保护需要同时考虑行为语义、模型能力和扰动机制，而不能只依赖单一准确率指标。",
    )

    add_heading(anchor, "4.6 第二阶段节点级功能测试结果分析", "Heading 2")
    add_body(
        anchor,
        "第一阶段实验关注数据侧扰动对行为序列攻击的影响，Cooja 节点级实验则把问题扩展到通信侧观察空间。节点程序通过 dummy_noise、dummy_ldp 与 dummy_adaptive_ldp 改变发送模式，"
        "日志解析流程再将应用层日志和 Radio log 转换为窗口级流量特征，使用 fixed_attacker 与 retrain_attacker 两种方式评估窃听者可见信息。",
    )
    add_body(
        anchor,
        "聚合结果显示，Cooja baseline_acc 为 0.2490。dummy_noise 在 fixed_attacker 下将 defended_acc 降至 0.1863，dummy_ldp 降至 0.2139，dummy_adaptive_ldp 降至 0.2096；"
        "retrain_attacker 下三者分别为 0.2598、0.2292 和 0.2427。结果说明节点侧 dummy 流量能够改变窃听者观察到的流量模式，但重训练攻击者仍可能利用部分残余结构。",
    )
    add_figure(anchor, FIGURES["cooja_accuracy"], "图4.8 Cooja 节点级 fixed_attacker/retrain_attacker 准确率对比")
    add_body(
        anchor,
        "需要强调的是，Cooja 部分用于节点侧功能性验证，不对真实能耗、真实端到端时延或 dummy/real 包比例作量化结论。"
        "当前日志可支持攻击准确率和窗口级功能验证，但 packet、byte、IAT 等字段在部分导出中不可用时，不应被解释为真实能耗或真实时延测量。",
    )

    add_heading(anchor, "4.7 真实数据扩展结果分析", "Heading 2")
    add_body(
        anchor,
        "真实数据扩展用于检验三类防御机制是否只在合成数据中有效。UCI HAR、Kasteren 和 CASAS 分别代表规则化人体活动窗口、细粒度智能家居事件和大规模房间级活动记录。"
        "最终真实数据主矩阵覆盖 3 个数据集、3 个 seed、2 类模型、3 种方法与 2 种威胁模式；参数扫描也覆盖三类方法的完整组合。",
    )
    add_body(
        anchor,
        "UCI HAR 的窗口形式规整、类别数较少，baseline 识别能力较高；Kasteren 的细粒度事件类别更多，识别任务更困难；CASAS hh101 样本规模较大，但活动类别和房间级行为模式更复杂。"
        "因此，不同数据集之间不宜直接比较绝对准确率，而应观察各自内部 baseline 到 defended 的变化。",
    )
    add_figure(
        anchor,
        FIGURES["real_accuracy"],
        "图4.9 真实数据集内部 baseline、fixed_attacker 与 retrain_attacker 准确率对比",
    )
    add_body(
        anchor,
        "由图4.9可见，三类真实数据在防御后均观察到准确率下降，说明防御效果并不只存在于 mock 场景。UCI HAR 展示了规则化传感器窗口中的攻击抑制现象，"
        "Kasteren 说明细粒度家居事件在扰动后更难识别，CASAS hh101 则支持大规模真实活动记录下的外部有效性。真实数据参数扫描已覆盖 UCI HAR、Kasteren 与 CASAS 三个数据集，"
        "不再局限于早期只展示 UCI HAR 的代表性口径。",
    )

    add_heading(anchor, "4.8 本章小结", "Heading 2")
    add_body(
        anchor,
        "本章基于最终完整实验结果，对 mock 主实验、真实公开数据扩展、参数扫描、混淆矩阵和 Cooja 节点级流量混淆进行了系统分析。"
        "攻击基线表明，LSTM 明显强于 MLP，说明物联网行为泄露更依赖时间结构和状态变化顺序，而不只是窗口统计量。",
    )
    add_body(
        anchor,
        "防御结果表明，默认参数下 LDP 对攻击抑制最强，但对应最高失真；noise 保留更多相关性，但剩余识别风险也较高；adaptive_ldp 在默认 profile 下体现了隐私—可用性折中，"
        "但面对不同模型和重训练攻击者时表现存在差异。参数扫描进一步说明隐私—可用性权衡是连续变化的，而 adaptive_ldp 的 6 个 profile 为预算范围、风险权重和边缘预算裁剪接口提供了经验性消融观察。",
    )
    add_body(
        anchor,
        "真实数据结果支持防御效果的外部有效性，Cooja 节点级实验则补充说明通信侧 dummy 流量能够改变窃听者可见模式。综合来看，物联网行为隐私保护需要同时考虑数据模态、攻击模型、扰动强度、攻击者适应能力和部署位置，"
        "不能把隐私保护简化为单一方法或单一指标的比较。",
    )


def verify_fonts(doc: Document) -> dict[str, object]:
    checked = 0
    missing_east_asia = 0
    wrong_latin = 0
    in_chapter4 = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "4 实验设计与结果分析" and paragraph.style.name == "Heading 1":
            in_chapter4 = True
            continue
        if text == "5 总结与展望" and paragraph.style.name == "Heading 1":
            in_chapter4 = False
        if not in_chapter4:
            continue
        if not paragraph.text.strip():
            continue
        if paragraph.style.name.startswith("Heading"):
            continue
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            checked += 1
            if run.font.name not in (None, "Times New Roman"):
                wrong_latin += 1
            rfonts = run._element.rPr.rFonts if run._element.rPr is not None and run._element.rPr.rFonts is not None else None
            east_asia = rfonts.get(qn("w:eastAsia")) if rfonts is not None else None
            if east_asia not in (None, "宋体"):
                missing_east_asia += 1
    return {
        "runs_checked": checked,
        "runs_with_non_times_new_romans": wrong_latin,
        "runs_with_non_songti_east_asia": missing_east_asia,
        "body_font_target": "中文宋体小四；英文和数字 Times New Roman 小四",
        "scope": "仅检查本次重写的第四章正文与图题",
    }


def report_path(path: Path) -> str:
    try:
        return path.resolve().relative_to(ROOT.resolve()).as_posix()
    except ValueError:
        return str(path)


def main() -> None:
    parser = argparse.ArgumentParser(description="重写论文第四章并插入基于最终结果生成的图。")
    parser.add_argument("--input", required=True, help="输入论文 docx")
    parser.add_argument("--output", default=str(ROOT / DEFAULT_OUTPUT_NAME), help="输出论文 docx")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    if not input_path.exists():
        raise FileNotFoundError(f"输入论文不存在：{input_path}")
    missing_figures = [str(path) for path in FIGURES.values() if not path.exists()]
    if missing_figures:
        raise FileNotFoundError(f"缺少论文图：{missing_figures}")

    doc = Document(str(input_path))
    replacements = replace_outdated_phrases(doc)
    anchor, removed_blocks = remove_chapter4(doc)
    insert_chapter4(anchor)
    add_update_fields_setting(doc)
    font_check = verify_fonts(doc)
    doc.save(str(output_path))

    report = {
        "input_docx": input_path.name,
        "output_docx": output_path.name,
        "output_location_note": "Word 文件输出到本地指定路径，未作为仓库实验产物提交。",
        "removed_chapter4_blocks": removed_blocks,
        "rewritten_sections": [
            "4 实验设计与结果分析",
            "4.1 实验环境与结果来源",
            "4.2 攻击基线结果",
            "4.3 多方法防御结果比较",
            "4.3.1 LSTM 结果分析",
            "4.3.2 MLP 结果分析",
            "4.3.3 参数扫描结果",
            "4.4 隐私—可用性权衡分析",
            "4.5 混淆矩阵与错误类型分析",
            "4.6 第二阶段节点级功能测试结果分析",
            "4.7 真实数据扩展结果分析",
            "4.8 本章小结",
        ],
        "figures_inserted": {key: report_path(path) for key, path in FIGURES.items()},
        "outdated_phrase_replacements": replacements,
        "training_rerun": False,
        "parameter_scan_rerun": False,
        "cooja_simulation_rerun": False,
        "toc_update": "已设置 updateFields；若 Word 未自动刷新，请打开文档后更新目录域。",
        "font_check": font_check,
    }
    (SUMMARY_DIR / "thesis_chapter4_rewrite_report.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    lines = [
        "# 第四章重写报告",
        "",
        f"- 输入 Word：`{input_path.name}`",
        f"- 输出 Word：`{output_path.name}`",
        "- 输出位置：本地指定路径，未作为仓库实验产物提交。",
        f"- 删除旧第四章 XML block 数：{removed_blocks}",
        "- 本次没有重跑训练实验、参数扫描或 Cooja 仿真。",
        "",
        "## 使用的 CSV/JSON 结果",
        "- `outputs/summaries/final_thesis/mock/mock_summary.csv`",
        "- `outputs/summaries/final_thesis/real/real_summary.csv`",
        "- `outputs/summaries/final_thesis/cooja/cooja_summary.csv`",
        "- `outputs/summaries/final_thesis/mock/mock_parameter_scan_ldp.csv`",
        "- `outputs/summaries/final_thesis/mock/mock_parameter_scan_noise.csv`",
        "- `outputs/summaries/final_thesis/mock/mock_parameter_scan_adaptive_ldp.csv`",
        "- `outputs/summaries/final_thesis/mock/mock_adaptive_ldp_ablation_summary.csv`",
        "- `outputs/experiments/mock/seed_42/**/confusion.json`",
        "",
        "## 插入或引用的论文图",
    ]
    lines.extend(f"- `{report_path(path)}`" for path in FIGURES.values())
    lines.extend(
        [
            "",
            "## 修正内容",
            "- 第四章已按最终完整实验矩阵重新组织。",
            "- 真实数据参数扫描表述已更新为覆盖 UCI HAR、Kasteren 与 CASAS。",
            "- adaptive_ldp 已加入 6-profile 级消融解释。",
            "- Cooja 部分只写 fixed/retrain 攻击准确率与节点侧功能性验证，不写真实能耗、真实端到端时延或 dummy/real 包比例量化。",
            "",
            "## 格式检查",
            f"- 字体目标：{font_check['body_font_target']}",
            f"- 检查正文 run 数：{font_check['runs_checked']}",
            f"- 非 Times New Roman run 数：{font_check['runs_with_non_times_new_romans']}",
            f"- 非宋体 eastAsia run 数：{font_check['runs_with_non_songti_east_asia']}",
            f"- 检查范围：{font_check['scope']}",
            "- 目录：已设置 Word 打开时更新域；如果本机 Word 未自动刷新，请打开后右键更新目录。",
        ]
    )
    (SUMMARY_DIR / "thesis_chapter4_rewrite_report.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(json.dumps({"output_docx": str(output_path), "report": str(SUMMARY_DIR / "thesis_chapter4_rewrite_report.json")}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
